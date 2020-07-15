import __init__
import logging
import os
from pathlib import Path
from unicodedata import normalize
import lxml.etree
import lxml.html
import parse, utils
import datetime
import const as CONST
import subprocess
import re
from docx.api import Document
import win32com.client
from win32com.client import Dispatch
import json
import time

logger = logging.getLogger(__name__)

class Base(object):
    def __init__(self, path):
        self.path = Path(path)

class FileT(Base):
    # Class FileT
    def __init__(self, path):
        super().__init__(path)
        self.doc = lxml.etree.parse(str(path))

    # Function get_tag: get the node of XML file with specific tag
    def get_tag(self, tag, index=0):
        '''Get normalized text of tag base on index of tag'''
        node = [e for e in self.doc.iterfind('.//{0}'.format(tag))][index]
        return node

    # Function update_tag: update the node with the specific tag and value
    def update_tag(self, tag, value):
        self.get_tag(tag).text = value
        path = re.sub("file:/", "", self.doc.docinfo.URL)
        with open(path, 'wb') as f:
            self.doc.write(f)

    # Function update_t: update t file with the input data
    def update_t(self, data):
        lst_header = ["UnitUnderTest", "ExecutionDate", "NTUserID", "FileName", "Verdict"]

        for h in lst_header:
            self.update_tag(h, data[h])

        # collect Percentage Coverage
        lst_Percentage = [e for e in self.doc.iterfind('.//{0}'.format("Percentage"))]

        lst_Percentage[0].text = data['C0']
        lst_Percentage[1].text = data['C1']
        if utils.load(CONST.SETTING).get("sheetname") == "Merged_J":
            lst_Percentage[2].text = data['MCDCU']

        path = re.sub("file:/", "", self.doc.docinfo.URL)
        with open(path, 'wb') as f:
            self.doc.write(f)

    # Function get_data: get the information in the Summary HTML file : "UnitUnderTest", "NTUserID", "FileName", "Verdict", "MetricName", "Percentage"
    def get_data(self):
        lst_header = ["UnitUnderTest", "NTUserID", "FileName", "Verdict", "MetricName", "Percentage"]
        ut = self.get_tag("UnitUnderTest").text
        user = self.get_tag("NTUserID").text
        func = self.get_tag("FileName").text
        result = self.get_tag("Verdict").text
        # collect title C0, C1, MC/DC
        lst_MetricName = [e.text for e in self.doc.iterfind('.//{0}'.format("MetricName"))]
        # collect Percentage Coverage
        lst_Percentage = [e.text for e in self.doc.iterfind('.//{0}'.format("Percentage"))]

        d = dict()
        lst_data = list()
        lst_data.append("NTUserID")
        lst_data.append(user)
        lst_data.append("FileName")
        lst_data.append(func)
        lst_data.append("TestResult")
        lst_data.append(result)

        for index, key in enumerate(lst_MetricName):
            lst_data.append(key)
            lst_data.append(lst_Percentage[index])

        # convert list to json
        d[ut] = {lst_data[i]: lst_data[i + 1] for i in range(0, len(lst_data), 2)}
        return d

class FileTestReportXML(Base):
    # Class FileTestReportXML
    def __init__(self, path):
        super().__init__(path)
        self.doc = lxml.etree.parse(str(path))

    # Function get_tag: get the node of XML file with specific tag
    def get_tag(self, tag, index=0):
        '''Get normalized text of tag base on index of tag'''
        node = [e for e in self.doc.iterfind('.//{0}'.format(tag))][index]
        return node

    # Function get_data: get the information in the Summary HTML file : Verdict, C0, C1, MCDC
    def get_data(self):
        lst_header = ["status", "statement", "decision", "booleanOperandEffectivenessMasking", "booleanOperandEffectivenessUnique", "testScriptName"]
        node_summary = self.get_tag("summary")
        status = {'Verdict': node_summary.attrib['status']}
        node_coverageInfo = self.get_tag("coverageInfo")[0]

        score = {'C0': item.text for item in node_coverageInfo if item.tag == "statement"}
        score = {**score, **{'C1': item.text for item in node_coverageInfo if item.tag == "decision"}}
        score = {**score, **{'MCDCM': item.text for item in node_coverageInfo if item.tag == "booleanOperandEffectivenessMasking"}}
        score = {**score, **{'MCDCU': item.text for item in node_coverageInfo if item.tag == "booleanOperandEffectivenessUnique"}}

        testscriptname = {'testScriptName': item.text for item in self.get_tag("info") if item.tag == "testScriptName"}

        data = dict()
        data = {**status, **score, **testscriptname}
        return data

class FileTestReportHTML(Base):
    # Class FileTestReportHTML
    def __init__(self, path):
        super().__init__(path)
        self.doc = lxml.html.parse(str(path))

    # Function get_tag: get the node of html file with specific tag
    def get_tag(self, tag, index=0):
        '''Get normalized text of tag base on index of tag'''
        node = [e for e in self.doc.iterfind('.//{0}'.format(tag))][index]
        return node

    # T.B.D
    def get_data(self):
        pass

class FileTestSummaryHTML(Base):
    # Class FileTestSummaryHTML
    def __init__(self, path):
        super().__init__(path)
        self.doc = lxml.html.parse(str(path))

    # Function get_tag: get the node of html file with specific tag
    def get_tag(self, tag, index=0):
        '''Get normalized text of tag base on index of tag'''
        node = [e for e in self.doc.iterfind('.//{0}'.format(tag))][index]
        return node

    # Function get_data: get the information in the Summary HTML file : Verdict, C0, C1, MCDC
    def get_data(self):
        data = dict()
        try:
            for e in self.doc.iterfind('.//{0}'.format("div")):
                if e.text == None:
                    continue
                if "Project:" in e.text or "Overall Result:" in e.text:
                    for item in e:
                        if "Project:" in e.text:
                            data = {**data, **{'Project': item.text}}
                        elif "Overall Result:" in e.text:
                            data = {**data, **{'Verdict': item.text}}
                        else:
                            print("BUG FileTestSummaryHTML get_data")

            key = ""
            flag = False
            for e in self.doc.iterfind('.//{0}'.format("div")):
                if e.text == None:
                    continue
                if e.text == "Statement (S)" or e.text == "Decision (D)" or e.text == "MC/DC - masking (M)" or e.text == "MC/DC - unique cause (U)":
                    flag = True
                    if e.text == "Statement (S)":
                        key = "C0"
                    elif e.text == "Decision (D)":
                        key = 'C1'
                    elif e.text == "MC/DC - masking (M)":
                        key = 'MCDCM'
                    elif e.text == "MC/DC - unique cause (U)":
                        key = 'MCDCU'
                    else:
                        raise("BUG")
                    next
                else:
                    if "%" in e.text and flag == True:
                        flag = False
                        val = e.text.replace("%", "")
                        data = {**data, key : val}
                        key = ""
                        val = ""
                        next
                    else:
                        continue

            flag = False
            for e in self.doc.iterfind('.//{0}'.format("div")):
                if e.text == None:
                    continue
                if e.text == "Summary generated":
                    flag = True
                    key = "date"
                    next
                else:
                    if flag == True:
                        flag = False
                        val = e.text
                        data = {**data, key : val}
                        break

        except Exception as e:
            data = {}
            raise(e)
        finally:
            return data

def reformat_string(value):
    temp = re.sub("[\n\t\r\x07\xa0]", " ", value.strip()).strip()
    temp = re.sub("\s+", " ", temp).strip()
    return temp

class FileWDoc(Base):
    def __init__(self, path):
        super().__init__(path)
        self.doc = str(path)

    # Function update_t: update t file with the input data
    def update(self, data, opt="TYPE_P"):
        if opt == "TYPE_P":
            document = Document(self.doc)
            table_infor = document.tables[1]
            table_infor.cell(0,1).text = data['date']
            table_infor.cell(0,3).text = data['project']
            table_infor.cell(0,5).text = data['review initiator']
            table_infor.cell(1,1).text = str(data['effort'])
            table_infor.cell(1,3).text = data['baseline']
            table_infor.cell(1,5).text = data['review partner']
            table_attach = document.tables[2]
            table_attach.cell(1, 2).text = data['path_testscript']
            table_attach.cell(3, 2).text = data['path_test_summary']
            table_attach.cell(3, 1).text = data['ScoreC0C1']
            document.save(self.doc)
        elif opt == "TYPE_A":
            pass
        else:
            raise "BUG No Type"

    # Function get_data: to get the array data with specfic key, value of the nested json
    def get_data(self, opt="TYPE_P"):
        dict_w = dict()
        if opt == "TYPE_P":
            document = Document(self.doc)
            table_infor = document.tables[1]
            table_attach = document.tables[2]
            table_finding = document.tables[3]
            table_check_list = document.tables[6]

            temp = re.sub("[\n\t]", " ", table_attach.cell(3, 1).text).strip()
            [score_c0, score_c1] = re.sub("^.*C0: ([0-9]+)%.*C1: ([0-9]+)%", r'\1 \2', temp).split(" ")

            finding = table_finding.cell(1, 1).text
            impact = table_finding.cell(1, 3).text
            confirm_UT9 = table_check_list.cell(12, 3).text

            dict_w = {
                'date': table_infor.cell(0,1).text,
                'project': table_infor.cell(0,3).text,
                'review initiator': table_infor.cell(0,5).text,
                'effort': table_infor.cell(1,1).text,
                'baseline': table_infor.cell(1,3).text,
                'review partner' : table_infor.cell(1,5).text,
                'path_testscript': table_attach.cell(1, 2).text,
                'path_test_summary': table_attach.cell(3, 2).text,
                'C0': score_c0,
                'C1': score_c1,
                'tbl_finding': {
                    "finding": finding,
                    "impact": impact,
                    "confirm_UT9": confirm_UT9
                }
            }

        elif opt == "TYPE_A":
            word = win32com.client.DispatchEx('Word.Application')
            word.Visible = 0
            word.DisplayAlerts = 0

            doc = word.Documents.Open(self.doc)
            table_infor = doc.Tables(2)
            table_attach = doc.Tables(3)
            table_finding = doc.Tables(4)
            table_check_list = doc.Tables(8)


            finding = reformat_string(table_finding.Cell(Row=2, Column=2).Range.Text)
            impact = reformat_string(table_finding.Cell(Row=2, Column=4).Range.Text)
            confirm_UT26 = reformat_string(table_check_list.Cell(Row=12, Column=5).Range.Text)

            temp = reformat_string(table_attach.Cell(Row=7, Column=2).Range.Text)
            [score_c0, score_c1, score_mcdc] = re.sub("^.*C0: ([0-9]+)%.*C1: ([0-9]+)%.*MCDC: ([0-9]+)%", r'\1 \2 \3', temp).split(" ")

            [project_name, item_revision] = re.sub("^(.*) v(.*)", r'\1 \2', reformat_string(table_infor.Cell(Row=1, Column=4).Range.Text)).split(" ")

            dict_w = {
                'date': reformat_string(table_infor.Cell(Row=1, Column=2).Range.Text),
                'project': project_name,
                'ItemRevision': item_revision.replace("_", "."),
                'review initiator': reformat_string(table_infor.Cell(Row=1, Column=6).Range.Text),
                'effort': reformat_string(table_infor.Cell(Row=2, Column=2).Range.Text),
                'baseline': reformat_string(table_infor.Cell(Row=2, Column=4).Range.Text),
                'review partner' : reformat_string(table_infor.Cell(Row=2, Column=6).Range.Text),
                'C0': score_c0,
                'C1': score_c1,
                'MCDC': score_mcdc,
                'tbl_finding': {
                    "finding": finding,
                    "impact": impact,
                    "confirm_UT26": confirm_UT26
                }
            }

            doc.Close()
            word.Quit()
            word = None
        else:
            dict_w = {}
            raise "BUG No Type"

        return dict_w

class FileCoverageReasonXLS(Base):
    def __init__(self, path):
        super().__init__(path)
        self.doc = str(path)

    # Function update_t: update t file with the input data
    def update(self, data):
        excel = win32com.client.Dispatch('Excel.Application')
        wb = excel.Workbooks.Open(self.doc)
        excel.Visible = False
        excel.DisplayAlerts = False
        wb.DoNotPromptForConvert = True
        wb.CheckCompatibility = False

        score_c0 = (value(formatNumber(float(value(data.get("C0"))) * 100)) if (value(data.get("C0")) != "-" and data.get("C0") != None) else "NA")
        score_c1 = (value(formatNumber(float(value(data.get("C1"))) * 100)) if (value(data.get("C1")) != "-" and data.get("C1") != None) else "NA")
        score_mcdc = (value(formatNumber(float(value(data.get("MCDC"))) * 100)) if (value(data.get("MCDC")) != "-" and data.get("MCDC") != None) else "NA")

        data_t = {
            "UnitUnderTest": data.get("ItemName"),
            "NTUserID": str(convert_name(key=data.get("Tester"), opt="id")),
            "ExecutionDate" : datetime.datetime.now().strftime("%Y-%m-%d"),
            "C0": score_c0,
            "C1": score_c1,
            "MCDCU": score_mcdc
        }

        writeData = wb.Worksheets(1)
        # Write data here
        infor_CoverageReasonXLS = utils.load(CONST.SETTING).get("CoverageReasonXLS")

        writeData.Range(infor_CoverageReasonXLS.get("Tester")).Value = data_t.get("NTUserID")
        writeData.Range(infor_CoverageReasonXLS.get("Date")).Value = data_t.get("ExecutionDate")
        writeData.Range(infor_CoverageReasonXLS.get("Item_Name")).Value = data_t.get("UnitUnderTest")
        writeData.Range(infor_CoverageReasonXLS.get("C0")).Value = data_t.get("C0")
        writeData.Range(infor_CoverageReasonXLS.get("C1")).Value = data_t.get("C1")
        writeData.Range(infor_CoverageReasonXLS.get("MCDC")).Value = data_t.get("MCDCU")

        wb.Save()
        wb.Close()
        excel.Quit()
        excel = None

    # Function get_data: to get the array data with specfic key, value of the nested json
    def get_data(self):
        excel = win32com.client.Dispatch('Excel.Application')
        wb = excel.Workbooks.Open(self.doc)

        excel.Visible = False
        excel.DisplayAlerts = False
        wb.DoNotPromptForConvert = True
        wb.CheckCompatibility = False

        readData = wb.Worksheets(1)
        allData = readData.UsedRange

        infor_CoverageReasonXLS = utils.load(CONST.SETTING).get("CoverageReasonXLS")

        data = {
            "Tester": value(allData.Cells(1, 2).value),
            "Date": value(allData.Cells(2, 2).value),
            "Item_Name": value(allData.Cells(3, 2).value),
            "C0": value(formatNumber(float(allData.Cells(9, 2).value))),
            "C1": value(formatNumber(float(allData.Cells(10, 2).value))),
            "MCDC": value(formatNumber(float(allData.Cells(11, 2).value)))
        }

        wb.Save()
        wb.Close()
        excel.Quit()
        excel = None
        return data

class FileSummaryXLSX(Base):
    # Class FileSummaryXLSX
    def __init__(self, path):
        super().__init__(path)
        self.doc = str(path)

    # Function parse2json: convert XLSX to json data
    def parse2json(self, sheetname="", begin=59, end=59):
        return dict(parse.parse_summary_json(self.doc, sheetname=sheetname, begin=begin, end=end))

    # Function get_data: to get the array data with specfic key, value of the nested json
    def get_data(self, data, key, value):
        item = -1
        d = dict()
        for item in data.keys():
            if(data[item].get(key) == value):
                d[item] = data[item]

        return d

class FileRCov(Base):
    # Class FileRCov
    def __init__(self, path):
        super().__init__(path)
        self.doc = str(path)

    # Function get_data: to get the array data with specfic key, value of the nested json
    def get_data(self):
        try:
            flag_count = 0

            data = dict()
            with open(self.doc, encoding='shift-jis', errors='ignore') as fp:
                for line in fp.readlines()[:100]:
                    line = line.strip()
                    if flag_count > 1:
                        break

                    if 'Conclusion :' in line:
                        flag_count = flag_count + 1
                        next
                    if (flag_count == 1):
                        if 'Statement blocks' in line or 'Decisions' in line or 'Modified conditions' in line:
                            temp = re.sub("\s+\.+\s+", ":", line)
                            temp = re.sub("%\s\(.*\)$", "", temp)
                            key = temp.split(":")[0]
                            val = convert_score_percentage(temp.split(":")[1], opt='nomul')
                            if key == "Statement blocks":
                                key = "C0"
                            elif key == "Decisions":
                                key = 'C1'
                            elif key == "Modified conditions":
                                key = 'MCDC'
                                rst = True

                            data = {**data, key : val}

        except:
            data = {}
        finally:
            return data

class FileTestDesignXLSX(Base):
    # Class FileTestDesignXLSX
    def __init__(self, path):
        super().__init__(path)
        self.doc = str(path)

    # Function parse2json: convert XLSX to json data
    def parse2json(self, sheetname="", begin=59, end=59):
        pass

    # Function get_data: to get the array data with specfic key, value of the nested json
    def get_data(self):
        if utils.load(CONST.SETTING).get("sheetname") == "Merged_J":
            item_name = re.sub("^TD_(\w.*)_(MT_\d.*)\.xls.*$", r'\1 \2', os.path.basename(self.doc)).split(" ")[0]
            item_revision = "NA"
        else:
            [item_name, item_revision] = re.sub("^TD_(\w.*)_v(\d.*)\.xls.*$", r'\1 \2', os.path.basename(self.doc)).split(" ")

        data = {
            "ItemName": item_name,
            "TM": parse.get_xlsx_cells(xlsx=self.doc, sheet="Testcases", list_cell=['A24']).get('A24'),
            "ItemRevision": item_revision.replace("_", "."),
            "Tester": parse.get_xlsx_cells(xlsx=self.doc, sheet="Revision History", list_cell=['C17']).get('C17'),
            "Date": parse.get_xlsx_cells(xlsx=self.doc, sheet="Revision History", list_cell=['D17']).get('D17'),
        }

        return data

class FileATTReportXML(Base):
    # Class FileTestReportXML
    def __init__(self, path):
        super().__init__(path)
        self.doc = lxml.etree.parse(str(path))

    # Function get_tag: get the node of XML file with specific tag
    def get_tag(self, tag, index=0):
        '''Get normalized text of tag base on index of tag'''
        node = [e for e in self.doc.iterfind('.//{0}'.format(tag))][index]
        return node

    # Function get_data: get the information in the Summary HTML file : Verdict, C0, C1, MCDC
    def get_data(self):
        lst_header = ["ClassName", "ClassVersion", "CompleteVerdict", "TestModuleName"]
        data = dict()

        for index, key in enumerate(lst_header):
            data = {**data, **{key: self.get_tag(key).text}}

        return data

# Check the directory of function is exist or not
def check_exist(dir_input, function):
    return Path(dir_input).joinpath(function).exists()

# Convert Tester name to Real Name or USERID
def convert_name(key, opt="name"):
    logger.debug("Convert name")
    try:
        users = utils.load(CONST.SETTING).get("users")

        if opt == "name" or opt == "id":
            return str(users[key].get(opt))
        else:
            raise("Bug convert name")
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Check path : have src extension or not
def is_have_src(path):
    return ("\\src" in path or "\\SRC" in path)

# Trim Component Path to find the correct path
def trim_src(path):
    result = re.sub("^.*\\\\rb\\\\", "rb\\\\", path)
    result = re.sub("^.*\\\\rba\\\\", "rba\\\\", result)
    result = re.sub("\\\\src\\\\.*$", "", result)
    result = re.sub("\\\\SRC\\\\.*$", "", result)
    result = re.sub("\\\\src$", "", result)
    result = re.sub("\\\\SRC$", "", result)
    return result

def value(cell):
    if str(cell).isdigit():
        return str(cell)
    else:
        return str(cell)

def convert_score_percentage(num, opt=""):
    if (utils.load(CONST.SETTING).get("sheetname") == "Merged_J"):
        if num == "NA" or num == "none":
            return "NA"

    if opt == 'nomul':
        return value(formatNumber(float(value(num))))
    else:
        return value(formatNumber(round(float(value(formatNumber(float(value(num)) * 100))), 1)))

def check_score(score_test_summary, score_exel, opt=""):
    return ((score_test_summary == convert_score_percentage(score_exel) if (value(score_exel) != "-" and score_exel != None) else "NA"))

def formatNumber(num):
    return int(num) if num % 1 == 0 else num

# Check the number attachment of OPL
def check_OPL_w(file):
    word = win32com.client.DispatchEx('Word.Application')
    word.Visible = 0
    word.DisplayAlerts = 0

    doc = word.Documents.Open(file.as_posix())

    num_OPL =  doc.InlineShapes.Count - 1

    doc.Close()
    word.Quit()
    word = None
    return num_OPL

# Check information between summary xlsx, and test_summay_html is same or not
def check_information(file_test_summary_html, data, function_with_prj_name="", file_test_report_xml="", file_t="", file_CoverageReasonXLS="", opt="", mode=""):
    try:
        data_test_summary = FileTestSummaryHTML(file_test_summary_html).get_data()

        logger.debug("Check information {}", data.get("ItemName").replace(".c", ""))
        count = 0
        flag = True

        sub_new_func = ""
        if (utils.load(CONST.SETTING).get("sheetname") == "Merged_J"):
            sub_new_func = function_with_prj_name
        else:
            sub_new_func = data.get("ItemName").replace(".c", "")

        if not (data_test_summary.get("Project") == sub_new_func):
            flag = False
            logger.error("Item FileTestSummaryHTML {} has different name: {} - {}".format(sub_new_func, data_test_summary.get("Project"), sub_new_func))

        if not (data_test_summary.get("Verdict") == "Pass"):
            flag = False
            logger.error("Item FileTestSummaryHTML {} got different Verdict: {} - {}".format(data_test_summary.get("Project"), data_test_summary.get("Verdict"), data.get("Status Result")))

        score_c0 = convert_score_percentage(data.get("C0"))
        score_c1 = convert_score_percentage(data.get("C1"))
        score_mcdc = convert_score_percentage(data.get("MCDC"))

        if not (check_score(score_test_summary=data_test_summary.get("C0"), score_exel=data.get("C0")) \
                and check_score(score_test_summary=data_test_summary.get("C1"), score_exel=data.get("C1")) \
                and check_score(score_test_summary=data_test_summary.get("MCDCU"), score_exel=data.get("MCDC"))):
            flag = False
            logger.error("Item FileTestSummaryHTML {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data_test_summary.get("Project"), data_test_summary.get("C0"), score_c0,
                                                                        data_test_summary.get("C1"), score_c1,
                                                                        data_test_summary.get("MCDCU"), score_mcdc)
                        )

        if (utils.load(CONST.SETTING).get("sheetname") == "Merged_J"):
            # Check information between FileT and Summary
            if (opt == "check_t_xls"):
                data_t = FileT(file_t).get_data()
                data_t = data_t[data.get("ItemName").replace(".c", "") + ".c"]
                if not (check_score(score_test_summary=data_t.get("C0"), score_exel=data.get("C0")) \
                        and check_score(score_test_summary=data_t.get("C1"), score_exel=data.get("C1")) \
                        and check_score(score_test_summary=data_t.get("MC/DC"), score_exel=data.get("MCDC"))):
                    flag = False
                    logger.error("Item FileT {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data_t.get("FileName").replace(".c", ""), data_test_summary.get("C0"), score_c0,
                                                                                data_t.get("C1"), score_c1,
                                                                                data_t.get("MC/DC"), score_mcdc)
                                )

            # Check information between FileTestReportXML and Summary
            data_test_report_xml = FileTestReportXML(file_test_report_xml).get_data()

            if not (check_score(score_test_summary=data_test_report_xml.get("C0"), score_exel=data.get("C0")) \
                    and check_score(score_test_summary=data_test_report_xml.get("C1"), score_exel=data.get("C1")) \
                    and check_score(score_test_summary=data_test_report_xml.get("MCDCU"), score_exel=data.get("MCDC"))):
                flag = False
                logger.error("Item FileTestReportXML {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data_test_report_xml.get("testScriptName"), data_test_report_xml.get("C0"), score_c0,
                                                                            data_test_report_xml.get("C1"), score_c1,
                                                                            data_test_report_xml.get("MCDCU"), score_mcdc)
                            )

            if (opt == "check_t_xls"):
                # Check information between FileCoverageReasonXLS and Summary
                data_CoverageReasonXLS = FileCoverageReasonXLS(file_CoverageReasonXLS).get_data()
                if not (check_score(score_test_summary=data_CoverageReasonXLS.get("C0"), score_exel=data.get("C0")) \
                        and check_score(score_test_summary=data_CoverageReasonXLS.get("C1"), score_exel=data.get("C1")) \
                        and check_score(score_test_summary=data_CoverageReasonXLS.get("MCDC"), score_exel=data.get("MCDC"))):
                    flag = False
                    logger.error("Item FileCoverageReasonXLS {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data_CoverageReasonXLS.get("Item_Name"), data_CoverageReasonXLS.get("C0"), score_c0,
                                                                                data_CoverageReasonXLS.get("C1"), score_c1,
                                                                                data_CoverageReasonXLS.get("MCDC"), score_mcdc)
                                )
        elif (utils.load(CONST.SETTING).get("sheetname") == "Merged_C"):
            if (opt == "check_WT"):
                file_w = utils.scan_files(Path(file_test_summary_html).parent.as_posix(), ext='.docx')[0][0]
                data_w = FileWDoc(file_Walkthrough).get_data()

                if not (data_w.get("project").replace(".c", "") == sub_new_func):
                    flag = False
                    logger.error("Item FileWDoc {} has wrong name - {}".format(data_w.get("project").replace(".c", ""), sub_new_func))

                if not (check_score(score_test_summary=data_w.get("C0"), score_exel=data.get("C0")) \
                        and check_score(score_test_summary=data_w.get("C1"), score_exel=data.get("C1"))):
                    flag = False
                    logger.error("Item FileWDoc {} has different C0: {}/{}; C1: {}/{}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("C0"), score_c0,
                                                                                data_w.get("C1"), score_c1))

                if data.get("Baseline") == "" or data.get("Baseline") == "None" or data.get("Baseline") == None:
                    temp_baseline = ""
                else:
                    temp_baseline = data.get("Baseline")

                if not ((data_w.get("baseline") == temp_baseline) or (temp_baseline == "" and data_Walkthrough.get("baseline") == "None")):
                    flag = False
                    logger.error("Item FileWDoc {} has different Baseline: {} - {}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("baseline"), temp_baseline)
                                    )

                if not (data_w.get("review partner") == convert_name(key=utils.load(CONST.SETTING, "users").get(data.get("Tester")).get("reviewer"), opt="name") \
                    and data_w.get("review initiator") == convert_name(key=data.get("Tester"), opt="name")):
                    flag = False

                    logger.error("Item FileWDoc has different reviewer/tester {}/{} - {}/{}".format(data_w.get("review partner"), convert_name(key=utils.load(CONST.SETTING, "users").get(data.get("Tester")).get("reviewer"), opt="name"), \
                                                                                  data_w.get("review initiator"), convert_name(key=data.get("Tester"), opt="name")
                                                                                 )
                                )

                temp_path_test_WT = str(trim_src(data.get("ComponentName"))) + "\\Unit_tst\\" + str(data.get("TaskID")) + "\\" + data.get("ItemName").replace(".c", "")
                path_testscript = temp_path_test_WT + "\\Test_Spec"
                path_test_summary = temp_path_test_WT + "\\Test_Result"

                if not(data_w.get('path_testscript') == path_testscript and data_Walkthrough.get('path_test_summary') == path_test_summary):
                    flag = False
                    logger.error("Item FileWDoc {} has wrong path: {}/{} - {}/{}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get('path_testscript'), path_testscript,
                                                                                data_w.get('path_test_summary'), path_test_summary)
                                )

                if (data.get("OPL/Defect") == "OPL" or data.get("OPL/Defect") == "Defect"):
                    num_OPL = check_OPL_w(file_Walkthrough)

                    if not (num_OPL > 0):
                        flag = False
                        logger.error("Item FileWDoc {} has none OPL: {}".format(data_w.get("project").replace(".c", ""), str(num_OPL)))

                    if not (data_w.get("tbl_finding").get('finding') != "/" \
                        and data_w.get("tbl_finding").get('impact') != "/"):
                        flag = False
                        logger.error("Item FileWDoc {} has none comment finding/impact: {} - {}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('finding'),
                                                                                    data_w.get("tbl_finding").get('impact'))
                                    )

                    if (data_w.get('C0') == "100" and data_Walkthrough.get('C1') == "100"):
                        if not (data_w.get("tbl_finding").get('confirm_UT9').strip() == "Yes, Documented"):
                            flag = False
                            logger.error("Item FileWDoc {} has wrong comment confirm UT9: {} - {}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT9'),
                                                                                        '"Yes, Documented"')
                                        )
                    else:
                        if not (data_w.get("tbl_finding").get('confirm_UT9').strip() == "No, Documented"):
                            flag = False
                            logger.error("Item FileWDoc {} has wrong comment confirm UT9: {} - {}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT9'),
                                                                                        '"No, Documented"')
                                    )
                else:
                    num_OPL = check_OPL_w(file_Walkthrough)

                    if not (num_OPL <= 0):
                        flag = False
                        logger.error("Item FileWDoc {} has OPL: {}".format(data_w.get("project").replace(".c", ""), str(num_OPL)))

                    if not (data_w.get("tbl_finding").get('finding') == "/" \
                        and data_w.get("tbl_finding").get('impact') == "/"):
                        flag = False
                        logger.error("Item FileWDoc {} has wrong comment finding/impact: {} - {}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('finding'),
                                                                                    data_w.get("tbl_finding").get('impact'))
                                    )

                    if not (data_w.get("tbl_finding").get('confirm_UT9').strip() == "Yes"):
                        flag = False
                        logger.error("Item FileWDoc {} has wrong comment confirm UT9: {} - {}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT9'),
                                                                                    '"Yes"')
                                    )

        else:
            raise ("No sheet name")

         if (datetime.datetime.strptime(data_test_summary.get("date"), "%b %d, %Y, %H:%M %p").strftime("%d-%b-%Y") == datetime.datetime.strptime(data.get("End"), "%Y-%m-%d %H:%M:%S").strftime("%d-%b-%Y")):
             flag = True
         else:
             flag = False
             logger.warning("ItemName {} got wrong date end: {} - {}".format(data_test_summary.get("Project"), datetime.datetime.strptime(data_test_summary.get("date"), "%b %d, %Y, %H:%M %p").strftime("%d-%b-%Y"), datetime.datetime.strptime(data.get("End"), "%Y-%m-%d %H:%M:%S").strftime("%d-%b-%Y")))
             return flag

        return flag
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

def check_information_TYPE_A(path, data, opt=""):
    def check_exist_plt(lst, pat):
        flag = False
        for f in lst:
            if f == pat:
                flag = True
                break

        return flag
    try:
        logger.debug("Check information {}", data.get("ItemName").replace(".c", ""))

        flag = True

        lst_file_test_design = utils.scan_files(Path(path).as_posix(), ext='.xlsm')[0]
        lst_file_R = utils.scan_files(Path(path).as_posix(), ext='ReportR.txt')[0]
        lst_file_ATT = utils.scan_files(Path(path).as_posix(), ext='ATT_Report.xml')[0]
        lst_file_PLT = [os.path.basename(f).replace('.plt', '') for f in utils.scan_files(Path(path).as_posix(), ext='.plt')[0] if re.search('mdfFiles', str(Path(f).as_posix)) == None]

        file_test_design = file_R = file_ATT = ""
        if not (len(lst_file_PLT)):
            flag = False
            logger.error("ItemName {} has none file PLT".format(data.get("ItemName").replace(".c", "")))
        else:
            lst_file_PLT = lst_file_PLT

        data_test_design = data_R = data_ATT = ""

        temp = ""
        if (utils.load(CONST.SETTING).get("sheetname") == "Merged_J"):
            temp = data.get("ItemName").replace(".c", "")
        else:
            temp = data.get("ItemName").replace(".c", "")

        if not (len(lst_file_test_design)):
            flag = False
            logger.error("ItemName {} has none file test design".format(data.get("ItemName").replace(".c", "")))
        else:
            file_test_design = lst_file_test_design[0]
            data_test_design = FileTestDesignXLSX(file_test_design).get_data()

            if not (len(lst_file_ATT)):
                flag = False
                logger.error("ItemName {} has none file ATT Report".format(data.get("ItemName").replace(".c", "")))
            else:
                file_ATT = lst_file_ATT[0]
                data_ATT = FileATTReportXML(file_ATT).get_data()

                if not (data_test_design.get("ItemName") == temp \
                    and data_ATT.get("ClassName") == temp):
                    flag = False
                    logger.error("Different ItemName {} - {}".format(data_test_design.get("ItemName"), temp))

                if utils.load(CONST.SETTING).get("sheetname") == "Merged_J":
                    if not(data_ATT.get("ClassVersion") == data.get("ItemRevision")):
                        flag = False
                        logger.error("ItemName {} got ItemRevision: {} - {}".format(data.get("ItemName"), data_ATT.get("ClassVersion"), data.get("ItemRevision")))
                else:
                    if not(data_ATT.get("ClassVersion") == data.get("ItemRevision") and data_test_design.get("ItemRevision") == data.get("ItemRevision")):
                        flag = False
                        logger.error("ItemName {} got ItemRevision: {}/{} - {}".format(data.get("ItemName"), data_ATT.get("ClassVersion"), data_test_design.get("ItemRevision"), data.get("ItemRevision")))

                # for tm_ATT in data_ATT.get("TestModuleName"):
                #     print(tm_ATT)

                if not check_exist_plt(lst_file_PLT, data_ATT.get("TestModuleName")):
                    flag = False
                    logger.error("ItemName {} got TestModuleName: {} - {}".format(data.get("ItemName"), data_ATT.get("TestModuleName"), lst_file_PLT))

                if not (data_ATT.get("CompleteVerdict") == "Passed"):
                    flag = False
                    logger.error("ItemName {} got Verdict: {} - {}".format(data.get("ItemName"), data_ATT.get("CompleteVerdict"), data.get("Status Result")))

        score_c0 = convert_score_percentage(data.get("C0"))
        score_c1 = convert_score_percentage(data.get("C1"))
        score_mcdc = convert_score_percentage(data.get("MCDC"))

        if not (len(lst_file_R)):
            flag = False
            logger.error("ItemName {} has none file R.txt".format(data.get("ItemName").replace(".c", "")))
        else:
            file_R = lst_file_R[0]
            data_R = FileRCov(file_R).get_data()
            if not (check_score(score_test_summary=data_R.get("C0"), score_exel=data.get("C0")) \
                    and check_score(score_test_summary=data_R.get("C1"), score_exel=data.get("C1")) \
                    and check_score(score_test_summary=data_R.get("MCDC"), score_exel=data.get("MCDC"))):
                flag = False
                logger.error("ItemName R Report {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data.get("ItemName"), data_R.get("C0"), score_c0,
                                                                            data_R.get("C1"), score_c1,
                                                                            data_R.get("MCDC"), score_mcdc)
                            )

        """ Check w"""
        if (utils.load(CONST.SETTING).get("sheetname") == "Merged_C"):
            if (opt == "check_WT"):
                lst_file_WT = utils.scan_files(Path(path).as_posix(), ext='.doc')[0]

                if not (len(lst_file_WT)):
                        flag = False
                        logger.error("ItemName {} has none file w".format(data.get("ItemName").replace(".c", "")))
                else:
                    file_w = lst_file_WT[0]
                    data_w = FileWDoc(file_Walkthrough).get_data(opt="TYPE_A")

                    if not (data_w.get("project").replace(".c", "") == data.get("ItemName")):
                        flag = False
                        logger.error("Item FileWDoc {} has wrong name - {}".format(data_w.get("project").replace(".c", ""), data.get("ItemName")))

                    if not(data_w.get("ItemRevision") == data.get("ItemRevision")):
                        flag = False
                        logger.error("Item FileWDoc {} has wrong ItemRevision: {} - {}".format(data.get("ItemName"), data_w.get("ItemRevision"), data.get("ItemRevision")))

                    if not (check_score(score_test_summary=data_w.get("C0"), score_exel=data.get("C0")) \
                            and check_score(score_test_summary=data_w.get("C1"), score_exel=data.get("C1"))):
                        flag = False
                        logger.error("Item FileWDoc {} has different C0: {}/{}; C1: {}/{}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("C0"), score_c0,
                                                                                    data_w.get("C1"), score_c1))

                    if data.get("Baseline") == "" or data.get("Baseline") == "None" or data.get("Baseline") == None:
                        temp_baseline = ""
                    else:
                        temp_baseline = data.get("Baseline")

                    if not ((data_w.get("baseline") == temp_baseline) or (temp_baseline == "" and data_Walkthrough.get("baseline") == "None")):
                        flag = False
                        logger.error("Item FileWDoc {} has different Baseline: {} - {}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("baseline"), temp_baseline)
                                        )

                    if not (data_w.get("review partner") == convert_name(key=utils.load(CONST.SETTING, "users").get(data.get("Tester")).get("reviewer"), opt="name") \
                        and data_w.get("review initiator") == convert_name(key=data.get("Tester"), opt="name")):
                        flag = False

                        logger.error("Item FileWDoc has different reviewer/tester {}/{} - {}/{}".format(data_w.get("review partner"), convert_name(key=utils.load(CONST.SETTING, "users").get(data.get("Tester")).get("reviewer"), opt="name"), \
                                                                                    data_w.get("review initiator"), convert_name(key=data.get("Tester"), opt="name")
                                                                                    )
                                    )

                    if (data.get("OPL/Defect") == "OPL" or data.get("OPL/Defect") == "Defect"):
                        num_OPL = check_OPL_w(file_Walkthrough)

                        if not (num_OPL > 0):
                            flag = False
                            logger.error("Item FileWDoc {} has none OPL: {}".format(data_w.get("project").replace(".c", ""), str(num_OPL)))

                        if not (data_w.get("tbl_finding").get('finding') != "/" \
                            and data_w.get("tbl_finding").get('impact') != "/"):
                            flag = False
                            logger.error("Item FileWDoc {} has none comment finding/impact: {} - {}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('finding'),
                                                                                        data_w.get("tbl_finding").get('impact'))
                                        )

                        if (data_w.get('C0') == "100" and data_Walkthrough.get('C1') == "100"):
                            if not (data_w.get("tbl_finding").get('confirm_UT26').strip() == "Yes, Documented"):
                                flag = False
                                logger.error("Item FileWDoc {} has wrong comment confirm UT26: {} - {}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT26'),
                                                                                            '"Yes, Documented"')
                                            )
                        else:
                            if not (data_w.get("tbl_finding").get('confirm_UT26').strip() == "No, Documented"):
                                flag = False
                                logger.error("Item FileWDoc {} has wrong comment confirm UT26: {} - {}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT26'),
                                                                                            '"No, Documented"')
                                        )
                    else:
                        num_OPL = check_OPL_w(file_Walkthrough)

                        if not (num_OPL <= 0):
                            flag = False
                            logger.error("Item FileWDoc {} has OPL: {}".format(data_w.get("project").replace(".c", ""), str(num_OPL)))

                        if not (data_w.get("tbl_finding").get('finding') == "/" \
                            and data_w.get("tbl_finding").get('impact') == "/"):
                            flag = False
                            logger.error("Item FileWDoc {} has wrong comment finding/impact: {} - {}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('finding'),
                                                                                        data_w.get("tbl_finding").get('impact'))
                                        )

                        if not (data_w.get("tbl_finding").get('confirm_UT26').strip() == ""):
                            flag = False
                            logger.error("Item FileWDoc {} has wrong comment confirm UT26: {} - {}".format(data_w.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT26'),
                                                                                        '""')
                                        )

        return flag
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Check Release for J is correct ot not
def check_archives_j(path_summary, dir_input, taskids, begin=59, end=59):
    logger.debug("Start checker: Archives")
    try:
        doc = FileSummaryXLSX(path_summary)
        data = doc.parse2json(begin=begin, end=end)

        file_log = open("log_delivery.txt", "w")

        print("Start checker: Archives")
        print("*****************************************************************")
        file_log.write("Start checker: Archives\n")
        file_log.write("*****************************************************************\n")

        for taskid in taskids["TaskGroup"]:
            temp_data_prj = doc.get_data(data=data, key="Project", value=taskids["Project"])
            data_taskid = doc.get_data(data=temp_data_prj, key="TaskGroup", value=taskid)
            bb_number = taskids["BB"]
            path_taskid = Path(dir_input).joinpath(str(taskid))
            if (path_taskid.exists()):
                count = 0
                for item in data_taskid.keys():

                    function = data_taskid[item].get("ItemName").replace(".c", "")
                    user_tester = data_taskid[item].get("Tester")
                    mt_number = data_taskid[item].get("MT_Number").replace("UT_", "").replace("MT_", "")

                    if "TYPE_A" == data_taskid[item].get("Type"):
                        mt_number = "MT_" + mt_number
                        bb_number = ""
                    elif "TYPE_P" == data_taskid[item].get("Type"):
                        mt_number = "UT_" + mt_number
                        bb_number = "_" + bb_number
                    else:
                        mt_number = "NONE"
                        raise "BUG mt_number"

                    folder_mt_function = "{}_{}{}".format(mt_number, function, bb_number)

                    b_check_exist = check_exist(dir_input=path_taskid, function=folder_mt_function)
                    if (b_check_exist):
                        count += 1

                        if "TYPE_A" == data_taskid[item].get("Type"):
                            if check_information_TYPE_A(path=Path.joinpath(path_taskid, folder_mt_function), data=data_taskid[item]):
                                print("{},{},{},{},{}".format(taskid, mt_number, function, user_tester, "OK"))
                                file_log.write("{},{},{},{},{}\n".format(taskid, mt_number, function, user_tester, "OK"))
                            else:
                                logger.error("Different Information {},{},{},{},{}".format(taskid, mt_number, function, user_tester, "NG_DiffInfor"))
                                file_log.write("{},{},{},{},{}\n".format(taskid, mt_number, function, user_tester, "NG_DiffInfor"))

                        elif "TYPE_P" == data_taskid[item].get("Type"):
                            file_t = Path(path_taskid).joinpath(folder_mt_function, "Cantata", "results", "{}.t".format(function))
                            file_CoverageReasonXLS = Path(path_taskid).joinpath(folder_mt_function, "doc", "{}_{}".format(function, "CodeCoverage_or_Fail_Reason.xls"))
                            file_test_report_xml = Path(path_taskid).joinpath(folder_mt_function, "Cantata", "results", "test_report.xml")
                            file_test_summary = Path(path_taskid).joinpath(folder_mt_function, "Cantata", "results", "test_summary.html")

                            option_check = ""
                            if (len(str(function)) < 32):
                                option_check = "check_t_xls"
                            else:
                                option_check = ""

                            if check_information(file_test_summary_html=file_test_summary, data=data_taskid[item], function_with_prj_name=folder_mt_function, file_test_report_xml=file_test_report_xml, file_t=file_t, file_CoverageReasonXLS=file_CoverageReasonXLS, opt=option_check):
                                print("{},{},{},{},{}".format(taskid, mt_number, function, user_tester, "OK"))
                                file_log.write("{},{},{},{},{}\n".format(taskid, mt_number, function, user_tester, "OK"))
                            else:
                                logger.error("Different Information {},{},{},{},{}".format(taskid, mt_number, function, user_tester, "NG_DiffInfor"))
                                file_log.write("{},{},{},{},{}\n".format(taskid, mt_number, function, user_tester, "NG_DiffInfor"))
                        else:
                            raise "Bug No Type"

                    else:
                        logging.warning("{},{},{},{},{}".format(taskid, mt_number, function, user_tester, "NG"))
                        file_log.write("{},{},{},{},{}\n".format(taskid, mt_number, function, user_tester, "NG"))

                num_t = len(utils.scan_files(directory=path_taskid, ext=".t")[0]) + len(utils.scan_files(directory=path_taskid, ext=".xlsm")[0])
                status = ["GOOD" if (num_t == len(data_taskid)) and (count == len(data_taskid)) else "BAD"][0]
                print("## Total {}: status {}: Found/Count/Total - {}/{}/{}".format(str(taskid), status, count, num_t, len(data_taskid)))
                print("-----------------------------------------------------------------\n")

                file_log.write("## Total {}: status {}: Found/Count/Total - {}/{}/{}\n".format(str(taskid), status, count, num_t, len(data_taskid)))
                file_log.write("-----------------------------------------------------------------\n")

            else:
                logger.warning("TaskID {} is not existed".format(path_taskid))
                file_log.write("TaskID {} is not existed\n".format(path_taskid))
                next

        print("FINISH")
        file_log.write("FINISH\n")
        file_log.close()

    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Check Release for J is correct ot not
def make_archives_j(path_summary, dir_input, dir_output, taskids, begin=59, end=59):
    logger.debug("Start checker: Make Archives J")
    try:
        doc = FileSummaryXLSX(path_summary)
        data = doc.parse2json(begin=begin, end=end)

        file_log = open("log_delivery.txt", "w")

        print("Start checker: Make Archives J")
        print("*****************************************************************")
        file_log.write("Start checker: Make Archives J\n")
        file_log.write("*****************************************************************\n")

        for taskid in taskids["TaskGroup"]:
            temp_data_prj = doc.get_data(data=data, key="Project", value=taskids["Project"])
            data_taskid = doc.get_data(data=temp_data_prj, key="TaskGroup", value=taskid)
            bb_number = taskids["BB"]
            path_taskid = Path(dir_input).joinpath(str(taskid))
            if (path_taskid.exists()):
                count = 0
                for item in data_taskid.keys():
                    function = data_taskid[item].get("ItemName").replace(".c", "")
                    user_tester = data_taskid[item].get("Tester")
                    mt_number = data_taskid[item].get("MT_Number").replace("UT_", "").replace("MT_", "")

                    if "TYPE_A" == data_taskid[item].get("Type"):
                        mt_number = "MT_" + mt_number
                        logging.warning("{},{},{},{}".format(taskid, function, user_tester, "NG_MT_Check_Later"))
                        file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "NG_MT_Check_Later"))
                        continue
                    elif "TYPE_P" == data_taskid[item].get("Type"):
                        mt_number = "UT_" + mt_number
                    else:
                        mt_number = "NONE"
                        print("BUG mt_number")

                    folder_mt_function = "{}_{}_{}".format(mt_number, function, bb_number)

                    b_check_exist = check_exist(dir_input=path_taskid, function=folder_mt_function)
                    if (b_check_exist):
                        count += 1

                        file_t = Path(path_taskid).joinpath(folder_mt_function, "Cantata", "results", "{}.t".format(function))
                        file_CoverageReasonXLS = Path(path_taskid).joinpath(folder_mt_function, "doc", "{}_{}".format(function, "CodeCoverage_or_Fail_Reason.xls"))
                        file_test_report_xml = Path(path_taskid).joinpath(folder_mt_function, "Cantata", "results", "test_report.xml")
                        file_test_summary = Path(path_taskid).joinpath(folder_mt_function, "Cantata", "results", "test_summary.html")

                        if check_information(file_test_summary_html=file_test_summary, data=data_taskid[item], function_with_prj_name=folder_mt_function, file_test_report_xml=file_test_report_xml, opt=""):
                            if (len(str(function)) < 32):
                                FileCoverageReasonXLS(file_CoverageReasonXLS).update(data_taskid[item])
                                utils.copy(src=Path(CONST.TEMPLATE).joinpath("template_j.t"), dst=file_t)
                                update_t(file=file_t, data=data_taskid[item], file_test_summary_html=file_test_summary)

                                print("{},{},{},{}".format(taskid, function, user_tester, "OK"))
                                file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "OK"))
                            else:
                                logger.error("Long Name {},{},{},{}".format(taskid, function, user_tester, "NG_Long_Name"))
                                file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "NG_Long_Name"))
                        else:
                            logger.error("Different Information {},{},{},{}".format(taskid, function, user_tester, "NG_DiffInfor"))
                            file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "NG_DiffInfor"))

                    else:
                        logging.warning("{},{},{},{}".format(taskid, function, user_tester, "NG"))
                        file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "NG"))

                num_t = len(utils.scan_files(directory=path_taskid, ext=".t")[0]) + len(utils.scan_files(directory=path_taskid, ext=".xlsm")[0])

                status = ["GOOD" if (num_t == len(data_taskid)) and (count == len(data_taskid)) else "BAD"][0]
                print("## Total {}: status {}: Found/Count/Total - {}/{}/{}".format(str(taskid), status, count, num_t, len(data_taskid)))
                print("-----------------------------------------------------------------\n")

                file_log.write("## Total {}: status {}: Found/Count/Total - {}/{}/{}\n".format(str(taskid), status, count, num_t, len(data_taskid)))
                file_log.write("-----------------------------------------------------------------\n")

            else:
                logger.warning("{} is not existed".format(path_taskid))
                file_log.write("{} is not existed\n".format(path_taskid))
                next

        print("FINISH")
        file_log.write("FINISH\n")
        file_log.close()

    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Check Release is correct or not
def check_releases(path_summary, dir_input, taskids, begin=59, end=59):
    logger.debug("Start checker: Release")

    try:
        doc = FileSummaryXLSX(path_summary)
        data = doc.parse2json(begin=begin, end=end)
        print("Start checker: Release")
        print("*****************************************************************")
        file_log = open("log_delivery.txt", "w")
        file_log.write("Start checker: Release\n")
        file_log.write("*****************************************************************\n")
        for taskid in taskids:
            data_taskid = doc.get_data(data=data, key="TaskID", value=taskid)
            path_taskid = Path(dir_input).joinpath(str(taskid), "RV")
            if (path_taskid.exists()):
                count = 0
                for item in data_taskid.keys():
                    function = data_taskid[item].get("ItemName").replace(".c", "")
                    user_tester = data_taskid[item].get("Tester")
                    b_check_exist = check_exist(dir_input=path_taskid, function=function)
                    if (b_check_exist):
                        count += 1
                        print("{},{},{},{}".format(taskid, function, user_tester, "OK"))
                        file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "OK"))

                    else:
                        logger.warning("{},{},{},{}".format(taskid, function, user_tester, "NG"))
                        file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "NG"))

                status = ["GOOD" if (len(os.listdir(path_taskid)) == len(data_taskid)) and (count == len(data_taskid)) else "BAD"][0]
                print("## Total {}: status {}: Found/Count/Total - {}/{}/{}".format(str(taskid), status, count, len(os.listdir(path_taskid)), len(data_taskid)))
                print("-----------------------------------------------------------------\n")

                file_log.write("## Total {}: status {}: Found/Count/Total - {}/{}/{}\n".format(str(taskid), status, count, len(os.listdir(path_taskid)), len(data_taskid)))
                file_log.write("-----------------------------------------------------------------\n")
            else:
                logger.error("{} is not existed".format(path_taskid))
                file_log.write("{} is not existed\n".format(path_taskid))
                next

        print("FINISH")
        file_log.write("FINISH\n")
        file_log.close()
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Check Archive is correct or not
def check_archives(path_summary, dir_input, taskids, begin=59, end=59):
    logger.debug("Start checker: Archives")
    try:
        doc = FileSummaryXLSX(path_summary)
        data = doc.parse2json(begin=begin, end=end)
        file_log = open("log_delivery.txt", "a")

        print("Start checker: Archives")
        print("*****************************************************************")
        file_log.write("Start checker: Archives\n")
        file_log.write("*****************************************************************\n")
        for taskid in taskids:
            data_taskid = doc.get_data(data=data, key="TaskID", value=taskid)
            path_taskid = Path(dir_input).joinpath(str(taskid), "AR")
            if (path_taskid.exists()):
                count = 0
                for item in data_taskid.keys():
                    function = data_taskid[item].get("ItemName").replace(".c", "")
                    user_tester = data_taskid[item].get("Tester")

                    path_with_component = ""
                    if "TYPE_A" == data_taskid[item].get("Type"):
                        path_with_component = Path(path_taskid)
                    elif "TYPE_P" == data_taskid[item].get("Type"):
                        path_with_component = Path(path_taskid).joinpath(str(trim_src(data_taskid[item].get("ComponentName"))), "Unit_tst", str(data_taskid[item].get("TaskID")))
                    else:
                        print("BUG No Type")

                    b_check_exist = check_exist(dir_input=path_with_component, function=function)
                    if (b_check_exist):
                        count += 1

                        if "TYPE_A" == data_taskid[item].get("Type"):
                            if check_information_TYPE_A(path=Path.joinpath(path_with_component, function), data=data_taskid[item], opt="check_WT"):
                                print("{},{},{},{}".format(taskid, function, user_tester, "OK"))
                                file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "OK"))
                            else:
                                logger.error("Different Information {},{},{},{}".format(taskid, function, user_tester, "NG_DiffInfor"))
                                file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "NG_DiffInfor"))
                        elif "TYPE_P" == data_taskid[item].get("Type"):
                            f_test_summary = Path(path_taskid).joinpath(path_with_component, function, "Test_Result", "test_summary.html")
                            if check_information(file_test_summary_html=f_test_summary, data=data_taskid[item], opt="check_WT"):
                                print("{},{},{},{}".format(taskid, function, user_tester, "OK"))
                                file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "OK"))
                            else:
                                logger.error("Different Information {},{},{},{}".format(taskid, function, user_tester, "NG_DiffInfor"))
                                file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "NG_DiffInfor"))
                        else:
                            raise "Bug: No type"

                    else:
                        logging.warning("{},{},{},{}".format(taskid, function, user_tester, "NG"))
                        file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "NG"))

                num_t = len(utils.scan_files(directory=path_taskid, ext=".t")[0])
                status = ["GOOD" if (num_t == len(data_taskid)) and (count == len(data_taskid)) else "BAD"][0]
                print("## Total {}: status {}: Found/Count/Total - {}/{}/{}".format(str(taskid), status, count, num_t, len(data_taskid)))
                print("-----------------------------------------------------------------\n")

                file_log.write("## Total {}: status {}: Found/Count/Total - {}/{}/{}\n".format(str(taskid), status, count, num_t, len(data_taskid)))
                file_log.write("-----------------------------------------------------------------\n")

            else:
                logger.warning("{} is not existed".format(path_taskid))
                file_log.write("{} is not existed\n".format(path_taskid))
                next

        print("FINISH")
        file_log.write("FINISH\n")
        file_log.close()

    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Update W
def update_w(file, data, file_test_summary_html):
    logger.debug("Update w {}", file)
    try:
        data_test_summary = FileTestSummaryHTML(file_test_summary_html).get_data()
        temp = str(trim_src(data.get("ComponentName"))) + "\\Unit_tst\\" + str(data.get("TaskID")) + "\\" + data.get("ItemName").replace(".c", "")

        data_baseline = data.get("Baseline")
        if data_baseline == None or data_baseline == "" or data_baseline == "None":
            data_baseline = ""

        reviewer = convert_name(key=utils.load(CONST.SETTING, "reviewer"), opt="name")

        dict_w = {
            'date': datetime.datetime.now().strftime("%m/%d/%Y"),
            'project': data.get("ItemName"),
            'review initiator': convert_name(key=data.get("Tester"), opt="name"),
            'effort': str(0.5),
            'baseline': data_baseline,
            'review partner' : reviewer,
            'path_testscript': temp + "\\Test_Spec",
            'path_test_summary': temp + "\\Test_Result",
            'ScoreC0C1': " Test summary\n\tC0: " + data_test_summary.get("C0") + "%\tC1: " + data_test_summary.get("C1") + "%",
        }

        FileWDoc(file).update(dict_w)

    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Update File T
def update_t(file, data, file_test_summary_html):
    logger.debug("Update t {}", file)

    try:
        data_test_summary = FileTestSummaryHTML(file_test_summary_html).get_data()
        data_t = {
            "UnitUnderTest": data.get("ItemName").replace(".c", "") + ".c",
            "NTUserID": str(convert_name(key=data.get("Tester"), opt="id")),
            "ExecutionDate" : datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ"),
            "FileName": data.get("ItemName"),
            "Verdict": data_test_summary.get('Verdict').replace("Pass", "Passed").replace("Fail", "Failed"),
            "C0": data_test_summary.get('C0'),
            "C1": data_test_summary.get('C1'),
            "MCDCU": data_test_summary.get('MCDCU'),
        }

        FileT(file).update_t(data_t)
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Encrypt file with 7z
def sevenzip(filename, zipname):
    zip_exe_path = utils.load(CONST.SETTING, "Tool_7z")
    with open(os.devnull, 'w') as null:
        subprocess.call([zip_exe_path, 'a', '-tzip', zipname, filename], stdout=null, stderr=null)

# Create Archive w
def make_archieves(path_summary, dir_input, dir_output, taskids, begin=59, end=59):
    logger.debug("Create Archive w")
    try:

        doc = FileSummaryXLSX(path_summary)
        data = doc.parse2json(begin=begin, end=end)

        print("Start checker: Make Archives")
        print("*****************************************************************")

        for taskid in taskids:
            data_taskid = doc.get_data(data=data, key="TaskID", value=taskid)
            path_taskid = Path(dir_input).joinpath(str(taskid), "RV")
            if (path_taskid.exists()):
                count = 0
                for item in data_taskid.keys():
                    function = data_taskid[item].get("ItemName").replace(".c", "")
                    user_tester = data_taskid[item].get("Tester")

                    b_check_exist = check_exist(dir_input=path_taskid, function=function)
                    if (b_check_exist):
                        count += 1
                        if "TYPE_A" == data_taskid[item].get("Type"):
                            logger.warning("{},{},{},{}".format(taskid, function, user_tester, "TYPE_A_No_Need"))
                            continue
                        elif "TYPE_P" == data_taskid[item].get("Type"):
                            temp = str(data_taskid[item].get("C0"))
                            if (temp != "-" or temp != ""):
                                temp_component = str(data_taskid[item].get("ComponentName"))
                                if is_have_src(temp_component):
                                    final_dst = Path(dir_output).joinpath(str(taskid), "AR", trim_src(temp_component), "Unit_tst", str(taskid), function)
                                    dir_Configuration = Path(final_dst).joinpath("Configuration")
                                    dir_Test_Spec = Path(final_dst).joinpath("Test_Spec")
                                    dir_Test_Result = Path(final_dst).joinpath("Test_Result")

                                    f_w = Path(dir_Test_Result).joinpath("Walkthrough_Protocol_" + function + ".docx")
                                    f_t = Path(dir_Test_Result).joinpath(function + ".t")
                                    f_test_summary = Path(path_taskid).joinpath(function, "Cantata", "results", "test_summary.html")

                                    if check_information(file_test_summary_html=f_test_summary, data=data_taskid[item]):
                                        Path(dir_Configuration).parent.mkdir(parents=True, exist_ok=True)
                                        Path(dir_Configuration).mkdir(exist_ok=True)
                                        Path(dir_Test_Spec).mkdir(exist_ok=True)
                                        Path(dir_Test_Result).mkdir(exist_ok=True)

                                        utils.copy(src=Path(CONST.TEMPLATE).joinpath("WT_template.docx"), dst=f_w)
                                        utils.copy(src=Path(CONST.TEMPLATE).joinpath("template.t"), dst=f_t)
                                        utils.copy(src=f_test_summary, dst=dir_Test_Result)

                                        update_w(file=f_walkthrough, data=data_taskid[item], file_test_summary_html=f_test_summary)
                                        update_t(file=f_t, data=data_taskid[item], file_test_summary_html=f_test_summary)

                                        sevenzip(filename=Path(path_taskid).joinpath(function).as_posix(), zipname=Path(dir_Configuration).joinpath(str(function) + ".zip").as_posix())

                                        for f in utils.scan_files(directory=Path(path_taskid).joinpath(function, "Cantata", "tests"), ext=".c")[0]:
                                            sevenzip(filename=f.as_posix(), zipname=Path(dir_Test_Spec).joinpath(os.path.basename(f).replace(".c", ".zip")).as_posix())

                                        print("{},{},{},{}".format(taskid, function, user_tester, "OK"))
                                    else:
                                        logger.error("Different information {},{},{},{}".format(taskid, function, user_tester, "NG"))
                                        next

                                else:
                                    logger.error("Miss src in componentname {},{},{},{}".format(taskid, function, user_tester, "NG"))
                        else:
                                raise "BUG No Type"
                    else:
                        logger.warning("{},{},{},{}".format(taskid, function, user_tester, "NG"))

                status = ["GOOD" if (len(os.listdir(path_taskid)) == len(data_taskid)) and (count == len(data_taskid)) else "BAD"][0]
                print("## Total {}: status {}: Found/Count/Total - {}/{}/{}".format(str(taskid), status, count, len(os.listdir(path_taskid)), len(data_taskid)))
                print("-----------------------------------------------------------------\n")
            else:
                next
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

def create_summary_json_file(file_summary, sheetname="", begin=59, end=59):
    # Generate json file
    if sheetname == "":
        sheetname = utils.load(CONST.SETTING).get("sheetname")

    file_log = Path(__file__).parent.joinpath("log_json", "{}_{}_{}.json".format("log_summary", sheetname, datetime.datetime.now().strftime("%Y_%m_%dT%H_%MZ"), ".json"))
    with open(file_log, errors='ignore', mode='w') as fp:
        json.dump(FileSummaryXLSX(file_summary).parse2json(sheetname=sheetname, begin=begin, end=end), fp, indent=4, sort_keys=True)

def collect_information_deliverables(file_summary, sheetname="", begin=59, end=59):
    if sheetname == "":
        sheetname = utils.load(CONST.SETTING).get("sheetname")

    data = FileSummaryXLSX(file_summary).parse2json(sheetname=sheetname, begin=begin, end=end)

    item = -1
    l_prj = list()
    for item in data.keys():
        if data[item].get("Project") is not None:
            l_prj.append(data[item].get("Project"))

    l_prj = list(set(l_prj))
    l_prj.sort()

    item = -1
    d = dict()

    for prj in l_prj:
        total_assign_asw = 0.0
        total_assign_psw = 0.0
        total_deliver_asw = 0.0
        total_deliver_psw = 0.0
        total_remain_asw = 0.0
        total_remain_psw = 0.0
        percentage_asw = 0.0
        percentage_psw = 0.0
        max_date_asw = ""
        max_date_psw = ""
        l_date_start_asw = list()
        l_date_end_asw = list()
        l_date_release_asw = list()
        l_date_start_psw = list()
        l_date_end_psw = list()
        l_date_release_psw = list()

        total_defect_asw = 0.0
        total_defect_psw = 0.0

        for item in data.keys():
            if(data[item].get("Project") == prj):
                if data[item].get("ELOC Recheck With Tool") is not None:
                    if "TYPE_A" == data[item].get("Type"):
                        total_assign_asw += float(data[item].get("ELOC Recheck With Tool"))
                    elif "TYPE_P" == data[item].get("Type"):
                        total_assign_psw += float(data[item].get("ELOC Recheck With Tool"))
                    else:
                        raise("BUG")
                else:
                    continue

                if data[item].get("LOC Complete") is not None:
                    if "TYPE_A" == data[item].get("Type"):
                        total_deliver_asw += float(data[item].get("LOC Complete"))
                    elif "TYPE_P" == data[item].get("Type"):
                        total_deliver_psw += float(data[item].get("LOC Complete"))
                    else:
                        raise("BUG")

                if data[item].get("Planned Start") is not None:
                    if "TYPE_A" == data[item].get("Type"):
                        l_date_start_asw.append(data[item].get("Planned Start"))
                    elif "TYPE_P" == data[item].get("Type"):
                        l_date_start_psw.append(data[item].get("Planned Start"))
                    else:
                        raise("BUG")

                if data[item].get("Planned End") is not None:
                    if "TYPE_A" == data[item].get("Type"):
                        l_date_end_asw.append(data[item].get("Planned End"))
                    elif "TYPE_P" == data[item].get("Type"):
                        l_date_end_psw.append(data[item].get("Planned End"))
                    else:
                        raise("BUG")

                if data[item].get("Release Date") is not None:
                    if "TYPE_A" == data[item].get("Type"):
                        l_date_release_asw.append(data[item].get("Release Date"))
                    elif "TYPE_P" == data[item].get("Type"):
                        l_date_release_psw.append(data[item].get("Release Date"))
                    else:
                        raise("BUG")

                if data[item].get("OPL/Defect") is not None:
                    if "TYPE_A" == data[item].get("Type"):
                        if "Defect" == data[item].get("OPL/Defect"):
                            total_defect_asw += 1
                    elif "TYPE_P" == data[item].get("Type"):
                        if "Defect" == data[item].get("OPL/Defect"):
                            total_defect_psw += 1
                    else:
                        raise("BUG")
                else:
                    continue


        date_start_asw = ""
        date_end_asw = ""
        date_release_asw = ""

        date_start_psw = ""
        date_end_psw = ""
        date_release_psw = ""

        if len(l_date_start_asw) == 0:
            date_start_asw = "NA"
        else:
            date_start_asw = min(l_date_start_asw)

        if len(l_date_end_asw) == 0:
            date_end_asw = "NA"
        else:
            date_end_asw = max(l_date_end_asw)

        if len(l_date_release_asw) == 0:
            date_release_asw = "NA"
        else:
            date_release_asw = max(l_date_release_asw)

        if len(l_date_start_psw) == 0:
            date_start_psw = "NA"
        else:
            date_start_psw = min(l_date_start_psw)

        if len(l_date_end_psw) == 0:
            date_end_psw = "NA"
        else:
            date_end_psw = max(l_date_end_psw)

        if len(l_date_release_psw) == 0:
            date_release_psw = "NA"
        else:
            date_release_psw = max(l_date_release_psw)

        total_remain_asw = total_assign_asw - total_deliver_asw
        total_remain_psw = total_assign_psw - total_deliver_psw

        if (total_assign_asw > 0):
            percentage_asw = round(total_deliver_asw/total_assign_asw * 100,2)
        elif (total_assign_asw == 0):
            percentage_asw = "NA"
        else:
            percentage_asw = "NG"

        if (total_assign_psw > 0):
            percentage_psw = round(total_deliver_psw/total_assign_psw * 100,2)
        elif (total_assign_psw == 0):
            percentage_psw = "NA"
        else:
            percentage_psw = "NG"

        template_json = {
            "Project": prj,
            "Type": "TYPE_A",
            "Assigned task (ELOC)": total_assign_asw,
            "Assigned date": date_start_asw,
            "Target date": date_end_asw,
            "Delivered task (ELOC)": total_deliver_asw,
            "Delivered date": date_release_asw,
            "Remain (ELOC)": total_remain_asw,
            "% Completion": percentage_asw
        }

        if total_assign_asw > 0:
            print("{},{},{},{},{},{},{},{},{},{}".format(prj, "TYPE_A", total_assign_asw, date_start_asw, date_end_asw, total_deliver_asw, date_release_asw, total_remain_asw, percentage_asw, total_defect_asw))
        if total_assign_psw > 0:
            print("{},{},{},{},{},{},{},{},{},{}".format(prj, "TYPE_P", total_assign_psw, date_start_psw, date_end_psw, total_deliver_psw, date_release_psw, total_remain_psw, percentage_psw, total_defect_psw))


    return l_prj

def make_folder_release(path_summary, l_packages, dir_output, begin=59, end=59):
    logger.debug("make_folder_release")
    try:
        doc = FileSummaryXLSX(path_summary)
        data = doc.parse2json(begin=begin, end=end)

        print("Start Making Folder Release")
        for package in l_packages:
            data_package = doc.get_data(data=data, key="Package", value=package)
            path_package = Path(dir_output).joinpath(str(package))
            Path(path_package).mkdir(exist_ok=True)
            if utils.load(CONST.SETTING).get("sheetname") == "Merged_C":
                path_package = Path(path_package).joinpath(str("C"))
            else:
                path_package = Path(path_package).joinpath(str("J"))
            Path(path_package).mkdir(exist_ok=True)

            for item in data_package.keys():
                taskid = data_package[item].get("TaskID")
                if taskid is not "None":
                    date_end = data_package[item].get("Planned End") if data_package[item].get("Planned End") != "None" else None
                    if date_end is not None:
                        date_end = datetime.datetime.strptime(date_end, "%Y-%m-%d %H:%M:%S").strftime("%d-%b-%Y")
                        dir_taskid = Path(path_package).joinpath(str(date_end), str(taskid))
                        dir_taskid_RV = Path(dir_taskid).joinpath(str("RV"))
                        dir_taskid_AR = Path(dir_taskid).joinpath(str("AR"))
                        Path(dir_taskid).parent.mkdir(parents=True, exist_ok=True)
                        Path(dir_taskid).mkdir(exist_ok=True)
                        Path(dir_taskid_RV).mkdir(exist_ok=True)
                        Path(dir_taskid_AR).mkdir(exist_ok=True)

                        print("{},{},{}".format(package, taskid,"OK"))
                    else:
                        logger.warning("Not filling Planned Start/Planned End: {},{},{}".format(package, taskid,"OK"))

        print("Finish Making Folder Release")
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Main
def main():
    try:
        print("<<Checker Version: {}>>".format(utils.load(CONST.VERSION).get("version")))

        file_summary = utils.load(CONST.SETTING).get("file_summary")
        sheetname = utils.load(CONST.SETTING).get("sheetname")
        dir_output = utils.load(CONST.SETTING).get("dir_output")

        l_taskids = utils.load(CONST.SETTING).get("l_taskids_c")
        dir_input = utils.load(CONST.SETTING).get("dir_input_c")

        lst_opt = utils.load(CONST.SETTING).get("mode_c")

        if sheetname == "Merged_J":
            l_taskids = utils.load(CONST.SETTING).get("l_taskids_j")
            dir_input = utils.load(CONST.SETTING).get("dir_input_j")
            lst_opt = utils.load(CONST.SETTING).get("mode_j")

        index_begin = utils.load(CONST.SETTING).get("coordinator").get("begin")
        index_end = utils.load(CONST.SETTING).get("coordinator").get("end")

        for opt in lst_opt:
            if opt == "check_releases":
                check_releases(path_summary=file_summary, dir_input=dir_input, taskids=l_taskids, begin=index_begin, end=index_end)
            elif opt == "check_archives":
                if sheetname == "Merged_J":
                    check_archives_j(path_summary=file_summary, dir_input=dir_input, taskids=l_taskids, begin=index_begin, end=index_end)
                else:
                    check_archives(path_summary=file_summary, dir_input=dir_input, taskids=l_taskids, begin=index_begin, end=index_end)
            elif opt == "make_archives":
                if sheetname == "Merged_J":
                    make_archives_j(path_summary=file_summary, dir_input=dir_input, dir_output=dir_output, taskids=l_taskids, begin=index_begin, end=index_end)
                else:
                    make_archieves(path_summary=file_summary, dir_input=dir_input, dir_output=dir_output, taskids=l_taskids, begin=index_begin, end=index_end)
            elif opt == "make_folder_release":
                """Make folder release"""
                l_folder_package = utils.load(CONST.SETTING).get("l_folder_package")
                make_folder_release(path_summary=file_summary, l_packages=l_folder_package, dir_output=dir_output, begin=index_begin, end=index_end)
            elif opt == "create_summary_json_file":
                """Create json file of summary to backup"""
                create_summary_json_file(file_summary=file_summary, sheetname="Merged_C", begin=index_begin, end=index_end)
                create_summary_json_file(file_summary=file_summary, sheetname="Merged_J", begin=index_begin, end=index_end)
            elif opt == "collect_information_deliverables":
                """Collect information for deliverables"""
                collect_information_deliverables(file_summary=file_summary, sheetname="Merged_J", begin=index_begin, end=index_end)
            else:
                raise("I dont know your mode")
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done)")

def check_update_version():
    directory="//hieu.nguyen-trung/script_auto_checker"
    data = utils.scan_files(directory, ext="version.json")[0]

    current_version = utils.load(CONST.VERSION).get("version")

    l_version = []
    for f in data:
        version = utils.load(f).get("version")
        l_version.append(version)

    l_version = list(set(l_version))
    l_version.sort()
    latest_version = ""
    if (len(l_version) > 0):
        latest_version = l_version[-1]
    else:
        raise "BUG check update version"

    flag = False

    print("-----------------------------------------------------------------")
    if latest_version != None and latest_version != current_version:
        print('Please checkout new version "{}" at: "{}"'.format(latest_version, directory))
        flag = True
    else:
        print("Your version is latest")
    print("-----------------------------------------------------------------")

    return flag

if __name__ == "__main__":
    if(check_update_version()):
        print("If you do not upgrade new version, you have to wait in 10s")
        for i in range(0, 11):
            print("------Waiting: {} s-------".format(str(9 - i)))
            time.sleep(1)

    main()
    os.system("pause")
